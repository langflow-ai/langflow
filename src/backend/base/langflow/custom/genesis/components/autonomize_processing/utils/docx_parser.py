"""DOCX Parser Utility Functions."""

import base64
import tempfile
from typing import Any, Dict, Union


def process_docx_file(file_content: str) -> Dict[str, Any]:
    """
    Process a DOCX file from base64 content and extract structured data.

    Args:
        file_content (str): Base64 encoded DOCX file content

    Returns:
        Dict[str, Any]: Structured document data with elements, images, and sections
    """
    try:
        # Decode base64 content
        if isinstance(file_content, str):
            try:
                file_bytes = base64.b64decode(file_content)
            except Exception as e:
                return {"error": f"Failed to decode base64 content: {str(e)}"}
        else:
            return {"error": "File content must be a base64 encoded string"}

        # Create a temporary file to work with
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
            tmp_file.write(file_bytes)
            tmp_file.flush()

            # Open document with python-docx
            try:
                from docx import Document
                doc = Document(tmp_file.name)
            except ImportError:
                return {"error": "python-docx is not available. Please install it with: pip install python-docx"}
            except Exception as e:
                return {"error": f"Failed to open DOCX document: {str(e)}"}

        # Extract structured data
        document_data = {
            "element_ids": [],
            "paragraphs": [],
            "tables": [],
            "images": [],
            "sections": {},
            "metadata": {
                "total_paragraphs": 0,
                "total_tables": 0,
                "total_images": 0
            }
        }

        element_counter = 0

        # Process paragraphs
        for para_idx, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():  # Only include non-empty paragraphs
                element_id = f"para_{element_counter}"
                element_counter += 1

                para_data = {
                    "element_id": element_id,
                    "text": paragraph.text.strip(),
                    "style": paragraph.style.name if paragraph.style else "Normal",
                    "type": "paragraph",
                    "index": para_idx
                }

                document_data["paragraphs"].append(para_data)
                document_data["element_ids"].append(element_id)

        # Process tables
        for table_idx, table in enumerate(doc.tables):
            element_id = f"table_{element_counter}"
            element_counter += 1

            table_data = {
                "element_id": element_id,
                "type": "table",
                "index": table_idx,
                "rows": []
            }

            for row_idx, row in enumerate(table.rows):
                row_data = {
                    "row_index": row_idx,
                    "cells": []
                }

                for cell_idx, cell in enumerate(row.cells):
                    cell_data = {
                        "cell_index": cell_idx,
                        "text": cell.text.strip()
                    }
                    row_data["cells"].append(cell_data)

                table_data["rows"].append(row_data)

            document_data["tables"].append(table_data)
            document_data["element_ids"].append(element_id)

        # Process images (basic image detection)
        try:
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    element_id = f"image_{element_counter}"
                    element_counter += 1

                    image_data = {
                        "element_id": element_id,
                        "type": "image",
                        "filename": rel.target_ref,
                        "content_type": getattr(rel, 'content_type', 'unknown')
                    }

                    document_data["images"].append(image_data)
                    document_data["element_ids"].append(element_id)
        except Exception:
            # If image processing fails, just continue without images
            pass

        # Process sections (headers and footers)
        try:
            for section_idx, section in enumerate(doc.sections):
                section_data = {
                    "index": section_idx,
                    "header": "",
                    "footer": ""
                }

                # Get header text
                if section.header:
                    header_text = []
                    for paragraph in section.header.paragraphs:
                        if paragraph.text.strip():
                            header_text.append(paragraph.text.strip())
                    section_data["header"] = "\n".join(header_text)

                # Get footer text
                if section.footer:
                    footer_text = []
                    for paragraph in section.footer.paragraphs:
                        if paragraph.text.strip():
                            footer_text.append(paragraph.text.strip())
                    section_data["footer"] = "\n".join(footer_text)

                document_data["sections"][f"section_{section_idx}"] = section_data
        except Exception:
            # If section processing fails, just continue without sections
            pass

        # Update metadata
        document_data["metadata"]["total_paragraphs"] = len(document_data["paragraphs"])
        document_data["metadata"]["total_tables"] = len(document_data["tables"])
        document_data["metadata"]["total_images"] = len(document_data["images"])
        document_data["metadata"]["total_elements"] = len(document_data["element_ids"])

        return document_data

    except Exception as e:
        return {"error": f"Unexpected error processing DOCX file: {str(e)}"}


def extract_text_from_docx(file_path: str) -> str:
    """
    Simple text extraction from DOCX file (fallback function).

    Args:
        file_path (str): Path to DOCX file

    Returns:
        str: Extracted text content
    """
    try:
        from docx import Document
        doc = Document(file_path)
        return "\n\n".join([paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()])
    except ImportError:
        raise ImportError("python-docx is not available. Please install it with: pip install python-docx")
    except Exception as e:
        raise ValueError(f"Error reading DOCX file: {str(e)}") from e