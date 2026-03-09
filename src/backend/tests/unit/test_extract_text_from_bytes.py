from io import BytesIO
from unittest.mock import MagicMock, patch

import pytest
from lfx.base.data.utils import extract_text_from_bytes
from pypdf import PdfWriter


def _make_blank_pdf(num_pages: int = 1) -> bytes:
    """Create a valid PDF with blank pages."""
    writer = PdfWriter()
    for _ in range(num_pages):
        writer.add_blank_page(width=612, height=792)
    buf = BytesIO()
    writer.write(buf)
    return buf.getvalue()


def _mock_pdf_reader(pages_text: list[str]):
    """Create a mock PdfReader that returns pages with given text."""
    mock_reader = MagicMock()
    mock_pages = []
    for text in pages_text:
        page = MagicMock()
        page.extract_text.return_value = text
        mock_pages.append(page)
    mock_reader.pages = mock_pages
    mock_reader.__enter__ = MagicMock(return_value=mock_reader)
    mock_reader.__exit__ = MagicMock(return_value=False)
    return mock_reader


class TestExtractTextFromBytesPDF:
    @patch("lfx.base.data.utils.PdfReader")
    def test_should_extract_text_from_valid_pdf(self, mock_reader_cls):
        mock_reader_cls.return_value = _mock_pdf_reader(["Hello World"])
        result = extract_text_from_bytes("document.pdf", _make_blank_pdf())
        assert "Hello World" in result

    @patch("lfx.base.data.utils.PdfReader")
    def test_should_extract_text_from_multi_page_pdf(self, mock_reader_cls):
        mock_reader_cls.return_value = _mock_pdf_reader(["Page one content", "Page two content"])
        result = extract_text_from_bytes("multi.pdf", _make_blank_pdf(2))
        assert "Page one content" in result
        assert "Page two content" in result

    @patch("lfx.base.data.utils.PdfReader")
    def test_should_join_pages_with_double_newline(self, mock_reader_cls):
        mock_reader_cls.return_value = _mock_pdf_reader(["First", "Second"])
        result = extract_text_from_bytes("test.pdf", _make_blank_pdf(2))
        assert result == "First\n\nSecond"

    @patch("lfx.base.data.utils.PdfReader")
    def test_should_be_case_insensitive_on_extension(self, mock_reader_cls):
        mock_reader_cls.return_value = _mock_pdf_reader(["Test"])
        result = extract_text_from_bytes("DOC.PDF", _make_blank_pdf())
        assert "Test" in result

    def test_should_raise_value_error_for_corrupted_pdf(self):
        with pytest.raises(ValueError, match="Failed to parse PDF file"):
            extract_text_from_bytes("bad.pdf", b"this is not a pdf")

    def test_should_raise_value_error_for_empty_pdf_bytes(self):
        with pytest.raises(ValueError, match="Failed to parse PDF file"):
            extract_text_from_bytes("empty.pdf", b"")

    def test_should_handle_pdf_with_blank_pages(self):
        result = extract_text_from_bytes("blank.pdf", _make_blank_pdf())
        assert isinstance(result, str)

    @patch("lfx.base.data.utils.PdfReader")
    def test_should_handle_page_returning_none(self, mock_reader_cls):
        mock_reader_cls.return_value = _mock_pdf_reader(["Text"])
        mock_reader_cls.return_value.pages[0].extract_text.return_value = None
        mock_reader_cls.return_value.__enter__.return_value = mock_reader_cls.return_value
        result = extract_text_from_bytes("null_page.pdf", _make_blank_pdf())
        assert isinstance(result, str)


class TestExtractTextFromBytesDOCX:
    def test_should_extract_text_from_valid_docx(self):
        from docx import Document

        doc = Document()
        doc.add_paragraph("Hello from DOCX")
        buf = BytesIO()
        doc.save(buf)

        result = extract_text_from_bytes("file.docx", buf.getvalue())
        assert "Hello from DOCX" in result

    def test_should_extract_multiple_paragraphs(self):
        from docx import Document

        doc = Document()
        doc.add_paragraph("First paragraph")
        doc.add_paragraph("Second paragraph")
        buf = BytesIO()
        doc.save(buf)

        result = extract_text_from_bytes("file.docx", buf.getvalue())
        assert "First paragraph" in result
        assert "Second paragraph" in result
        assert "\n\n" in result

    def test_should_be_case_insensitive_on_extension(self):
        from docx import Document

        doc = Document()
        doc.add_paragraph("Case test")
        buf = BytesIO()
        doc.save(buf)

        result = extract_text_from_bytes("FILE.DOCX", buf.getvalue())
        assert "Case test" in result

    def test_should_raise_value_error_for_corrupted_docx(self):
        with pytest.raises(ValueError, match="Failed to parse DOCX file"):
            extract_text_from_bytes("bad.docx", b"not a valid docx")

    def test_should_raise_value_error_for_empty_docx_bytes(self):
        with pytest.raises(ValueError, match="Failed to parse DOCX file"):
            extract_text_from_bytes("empty.docx", b"")

    def test_should_handle_docx_with_no_paragraphs(self):
        from docx import Document

        doc = Document()
        buf = BytesIO()
        doc.save(buf)

        result = extract_text_from_bytes("empty_doc.docx", buf.getvalue())
        assert isinstance(result, str)


class TestExtractTextFromBytesPlainText:
    def test_should_decode_utf8_text(self):
        content = b"Hello plain text"
        result = extract_text_from_bytes("readme.txt", content)
        assert result == "Hello plain text"

    def test_should_handle_non_utf8_gracefully(self):
        content = b"\xff\xfe\x00\x01 some text"
        result = extract_text_from_bytes("binary.txt", content)
        assert isinstance(result, str)
        assert "some text" in result

    def test_should_handle_empty_content(self):
        result = extract_text_from_bytes("empty.txt", b"")
        assert result == ""

    def test_should_handle_csv_as_plain_text(self):
        content = b"col1,col2\nval1,val2"
        result = extract_text_from_bytes("data.csv", content)
        assert "col1,col2" in result

    def test_should_handle_json_as_plain_text(self):
        content = b'{"key": "value"}'
        result = extract_text_from_bytes("data.json", content)
        assert '"key"' in result

    def test_should_handle_unknown_extension_as_plain_text(self):
        content = b"some content"
        result = extract_text_from_bytes("file.xyz", content)
        assert result == "some content"

    def test_should_handle_file_without_extension(self):
        content = b"no extension"
        result = extract_text_from_bytes("Makefile", content)
        assert result == "no extension"

    def test_should_preserve_unicode_characters(self):
        content = "café résumé naïve".encode()
        result = extract_text_from_bytes("unicode.txt", content)
        assert result == "café résumé naïve"
