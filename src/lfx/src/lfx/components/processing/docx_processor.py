"""DOCX Processor Component - Simple wrapper for DOCX extraction."""

import base64

import aiohttp
from lfx.custom.custom_component.component import Component
from langflow.inputs import DataInput
from lfx.io import BoolInput, MessageTextInput, Output
from lfx.schema.data import Data

# Inline utility function
import base64
import tempfile


def process_docx_file(file_content: str) -> dict:
    """Process a DOCX file from base64 content and extract structured data."""
    try:
        # Decode base64 content
        if isinstance(file_content, str):
            try:
                file_bytes = base64.b64decode(file_content)
            except Exception as e:
                return {"error": f"Failed to decode base64 content: {str(e)}"}
        else:
            return {"error": "File content must be a base64 encoded string"}

        # Create a temporary file to work with
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
            tmp_file.write(file_bytes)
            tmp_file.flush()

            # Open document with python-docx
            try:
                from docx import Document
                doc = Document(tmp_file.name)
            except ImportError:
                return {"error": "python-docx is not available. Please install it with: pip install python-docx"}
            except Exception as e:
                return {"error": f"Failed to open DOCX document: {str(e)}"}

        # Extract structured data
        document_data = {
            "element_ids": [],
            "paragraphs": [],
            "tables": [],
            "images": [],
            "sections": {},
            "metadata": {
                "total_paragraphs": 0,
                "total_tables": 0,
                "total_images": 0
            }
        }

        element_counter = 0

        # Process paragraphs
        for para_idx, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():  # Only include non-empty paragraphs
                element_id = f"para_{element_counter}"
                element_counter += 1

                para_data = {
                    "element_id": element_id,
                    "text": paragraph.text.strip(),
                    "style": paragraph.style.name if paragraph.style else "Normal",
                    "type": "paragraph",
                    "index": para_idx
                }

                document_data["paragraphs"].append(para_data)
                document_data["element_ids"].append(element_id)

        # Process tables
        for table_idx, table in enumerate(doc.tables):
            element_id = f"table_{element_counter}"
            element_counter += 1

            table_data = {
                "element_id": element_id,
                "type": "table",
                "index": table_idx,
                "rows": []
            }

            for row_idx, row in enumerate(table.rows):
                row_data = {
                    "row_index": row_idx,
                    "cells": []
                }

                for cell_idx, cell in enumerate(row.cells):
                    cell_data = {
                        "cell_index": cell_idx,
                        "text": cell.text.strip()
                    }
                    row_data["cells"].append(cell_data)

                table_data["rows"].append(row_data)

            document_data["tables"].append(table_data)
            document_data["element_ids"].append(element_id)

        # Process images (basic image detection)
        try:
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    element_id = f"image_{element_counter}"
                    element_counter += 1

                    image_data = {
                        "element_id": element_id,
                        "type": "image",
                        "filename": rel.target_ref,
                        "content_type": getattr(rel, 'content_type', 'unknown')
                    }

                    document_data["images"].append(image_data)
                    document_data["element_ids"].append(element_id)
        except Exception:
            # If image processing fails, just continue without images
            pass

        # Process sections (headers and footers)
        try:
            for section_idx, section in enumerate(doc.sections):
                section_data = {
                    "index": section_idx,
                    "header": "",
                    "footer": ""
                }

                # Get header text
                if section.header:
                    header_text = []
                    for paragraph in section.header.paragraphs:
                        if paragraph.text.strip():
                            header_text.append(paragraph.text.strip())
                    section_data["header"] = "\n".join(header_text)

                # Get footer text
                if section.footer:
                    footer_text = []
                    for paragraph in section.footer.paragraphs:
                        if paragraph.text.strip():
                            footer_text.append(paragraph.text.strip())
                    section_data["footer"] = "\n".join(footer_text)

                document_data["sections"][f"section_{section_idx}"] = section_data
        except Exception:
            # If section processing fails, just continue without sections
            pass

        # Update metadata
        document_data["metadata"]["total_paragraphs"] = len(document_data["paragraphs"])
        document_data["metadata"]["total_tables"] = len(document_data["tables"])
        document_data["metadata"]["total_images"] = len(document_data["images"])
        document_data["metadata"]["total_elements"] = len(document_data["element_ids"])

        return document_data

    except Exception as e:
        return {"error": f"Unexpected error processing DOCX file: {str(e)}"}


class DocxProcessorComponent(Component):
    """DOCX Processor Component for extracting content from DOCX files."""

    display_name = "DOCX Processor"
    category: str = "processing"
    description = "Extract structured content from DOCX files"
    documentation = ""
    icon = "file-text"
    name = "DocxProcessor"

    inputs = [
        DataInput(
            name="file_input",
            display_name="File Input",
            info="Data object with file_path (signed URL) or base64 content",
            required=False,
            is_list=True,  # BlobStorage returns list[Data]
        ),
        MessageTextInput(
            name="direct_input",
            display_name="Direct Input (Optional)",
            info="Base64 encoded DOCX content or signed URL (if not using Data input)",
            value="",
            required=False,
        ),
        BoolInput(
            name="extract_images",
            display_name="Extract Images",
            value=True,
            info="Extract and encode embedded images",
        ),
        BoolInput(
            name="extract_sections",
            display_name="Extract Sections",
            value=True,
            info="Extract header and footer sections",
        ),
    ]

    outputs = [
        Output(display_name="Document", name="document", method="process_document"),
    ]

    async def process_document(self) -> Data:
        """Process the DOCX document and return extracted data."""
        try:
            # Get the file content - either from Data input or direct input
            file_content = None

            # First check Data input (from BlobStorage)
            if self.file_input:
                # Debug: Log what we received
                self.status = f"Received input type: {type(self.file_input)}"

                # Handle list of Data objects (BlobStorage returns list[Data])
                if isinstance(self.file_input, list):
                    if len(self.file_input) > 0:
                        data_obj = self.file_input[0]  # Process first file
                        self.status = (
                            f"Processing first file from list of {len(self.file_input)}"
                        )
                    else:
                        return Data(value={"error": "Empty file list provided"})
                else:
                    data_obj = self.file_input

                # Handle Data object
                if hasattr(data_obj, "data"):
                    file_data = data_obj.data
                elif hasattr(data_obj, "value"):
                    file_data = data_obj.value
                else:
                    file_data = data_obj

                # Extract file_path (signed URL) from Data object
                if isinstance(file_data, dict) and "file_path" in file_data:
                    signed_url = file_data["file_path"]
                    self.status = f"Downloading from URL: {signed_url[:50]}..."

                    # Download file from signed URL
                    async with aiohttp.ClientSession() as session:
                        async with session.get(signed_url) as response:
                            if response.status == 200:
                                file_bytes = await response.read()
                                # Convert to base64 for processing
                                file_content = base64.b64encode(file_bytes).decode(
                                    "utf-8"
                                )
                                self.status = f"Downloaded {len(file_bytes)} bytes"
                            else:
                                return Data(
                                    value={
                                        "error": f"Failed to download file: {response.status}"
                                    }
                                )
                elif isinstance(file_data, str):
                    # Direct base64 content
                    file_content = file_data
                    self.status = "Using direct base64 content"
                else:
                    return Data(
                        value={
                            "error": f"Unexpected data format: {type(file_data)}, content: {str(file_data)[:100]}"
                        }
                    )

            # Fallback to direct input if no Data input
            if not file_content and self.direct_input:
                file_content = self.direct_input

            if not file_content:
                return Data(value={"error": "No file input provided"})

            # Process the document using utility function
            document = process_docx_file(file_content)

            # Check for errors
            if "error" in document:
                self.status = f"Error: {document['error']}"
                return Data(value=document)

            # Filter based on options
            if not self.extract_images:
                document.pop("images", None)

            if not self.extract_sections:
                document.pop("sections", None)

            # Set status
            element_count = len(document.get("element_ids", []))
            self.status = f"Processed document with {element_count} elements"

            return Data(value=document)

        except Exception as e:
            self.status = f"Error: {str(e)}"
            return Data(value={"error": str(e), "status": "failed"})
